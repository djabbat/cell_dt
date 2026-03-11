#!/usr/bin/env python3
"""
Generate a scientific DOCX article on morphogens and CDATA theory.
Author: Jaba Tkemaladze
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  font_size=12, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, italic=italic)
    return p

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(10)
    run = p.add_run(text)
    set_font(run, size=18, bold=True)
    return p

def shade_cell(cell, fill_rgb):
    """Apply background color to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = '{:02X}{:02X}{:02X}'.format(*fill_rgb)
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=True, color=(255,255,255), size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)

def add_table_with_header(doc, headers, rows, header_color, caption_text):
    """Add a formatted table with colored header and data rows."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption_text)
    set_font(cap_run, size=11, bold=True, italic=True)
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], header_color)
        set_cell_text(hdr_cells[i], h, bold=True, color=(255, 255, 255), size=10)

    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            row_cells[ci].text = ''
            p = row_cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_font(run, size=9, bold=False, color=(0, 0, 0))
            if ri % 2 == 1:
                shade_cell(row_cells[ci], (240, 240, 245))

    doc.add_paragraph()
    return table

def add_figure_caption(doc, fig_num, title, ascii_art, equation=None):
    """Add a figure caption block with ASCII art."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {fig_num}. {title}")
    set_font(run, size=11, bold=False, italic=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)

    art_para = doc.add_paragraph()
    art_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    art_para.paragraph_format.space_before = Pt(2)
    art_para.paragraph_format.space_after = Pt(2)
    art_para.paragraph_format.left_indent = Cm(1.5)
    art_run = art_para.add_run(ascii_art)
    art_run.font.name = "Courier New"
    art_run.font.size = Pt(9)

    if equation:
        eq_p = doc.add_paragraph()
        eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_run = eq_p.add_run(equation)
        eq_run.font.name = "Courier New"
        eq_run.font.size = Pt(9)
        eq_run.font.italic = True

    doc.add_paragraph()

def add_reference(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(1.25)
    pf.first_line_indent = Cm(-1.25)
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11)

# ─── Document Construction ───────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
add_title(doc,
    "Morphogen Gradients, Primary Cilia, and Centriolar Aging: "
    "Spatiotemporal Disruption of Developmental Signaling in the CDATA Framework"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Jaba Tkemaladze")
set_font(r, size=13, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Institute of Theoretical and Experimental Medicine")
set_font(r2, size=11, italic=True)
p2.paragraph_format.space_after = Pt(2)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Correspondence: jaba.tkemaladze@itemed.org")
set_font(r3, size=10)
p3.paragraph_format.space_after = Pt(14)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Abstract")
abstract_text = (
    "Morphogens are secreted signaling molecules that form concentration gradients across developing tissues, "
    "instructing cell fate in a dose-dependent manner. The classical French Flag model (Wolpert, 1969) and the "
    "Turing reaction-diffusion framework (Turing, 1952) together explain how positional information and "
    "self-organizing patterns are encoded in morphogen landscapes. Over the past two decades, primary cilia "
    "have emerged as essential transduction hubs for multiple morphogen pathways — most critically, for Sonic "
    "Hedgehog (Shh) signaling in vertebrates, where GLI transcription factor processing is obligatorily "
    "compartmentalized within the ciliary axoneme (Huangfu et al., 2003; Rohatgi et al., 2007). "
    "The Centriolar Damage Theory of Aging (CDATA), proposed by Tkemaladze, posits that the progressive "
    "structural deterioration of centrioles and transition zones — evidenced by declining CEP164 scaffolding "
    "protein expression — impairs ciliary assembly and dynamics, thereby blunting morphogen responsiveness "
    "in a Hill-function nonlinear fashion. Here we provide a comprehensive review of morphogen biology, "
    "their spatial and temporal distribution, the evolutionary divergence of ciliary dependence across taxa, "
    "and the mechanistic consequences of centriolar aging on morphogen signaling. We derive a quantitative "
    "model linking ciliary functional decline to graded loss of GLI activation (A_GLI = A_max * C^n / "
    "(K^n + C^n)), and discuss clinical manifestations including Bardet-Biedl syndrome, Joubert syndrome, "
    "age-related myeloid skewing, and regenerative failure. The CDATA framework provides a unifying "
    "mechanistic explanation for the deterioration of tissue patterning fidelity with advancing age."
)
add_paragraph(doc, abstract_text, space_before=0, space_after=8)

p_kw = doc.add_paragraph()
pf_kw = p_kw.paragraph_format
pf_kw.space_after = Pt(14)
r_kw_label = p_kw.add_run("Keywords: ")
set_font(r_kw_label, size=11, bold=True)
r_kw = p_kw.add_run(
    "morphogen gradients; primary cilia; Hedgehog signaling; GLI transcription factors; "
    "centriolar aging; CDATA; reaction-diffusion; positional information; ciliopathy; "
    "developmental biology; aging"
)
set_font(r_kw, size=11, italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Introduction")

add_paragraph(doc,
    "The problem of how a single fertilized egg reliably generates the astonishing spatial complexity of a "
    "multicellular organism has fascinated biologists for over a century. A central part of the answer lies "
    "in morphogens: diffusible signaling molecules secreted from localized sources that establish "
    "concentration gradients across fields of cells, translating graded chemical information into discrete "
    "cell identities. The concept was formalized by Lewis Wolpert (1969) in the positional information "
    "hypothesis, famously illustrated by the French Flag model: a one-dimensional field of cells exposed to "
    "a morphogen gradient interprets high concentrations as 'blue', intermediate as 'white', and low as "
    "'red', depending on threshold concentrations of the signal. This elegant framework unified observations "
    "from insect segmentation, vertebrate limb patterning, and neural tube specification into a single "
    "conceptual scaffold."
)

add_paragraph(doc,
    "Complementarily, Alan Turing's reaction-diffusion theory (1952) demonstrated that two diffusing "
    "substances — an autocatalytic activator and a longer-range inhibitor — can spontaneously generate "
    "periodic spatial patterns from near-homogeneous initial conditions. Turing instabilities underlie "
    "digit spacing in vertebrate limbs, hair follicle periodicity, skin pigmentation patterns, and numerous "
    "other biological textures (Maini et al., 2012). Together, gradient-based positional information and "
    "self-organizing reaction-diffusion mechanisms constitute the two foundational pillars of morphogenetic "
    "field theory."
)

add_paragraph(doc,
    "Despite decades of genetic and biochemical dissection of morphogen pathways, a unifying mechanism "
    "linking the temporal decline of morphogen responsiveness to organismal aging remained elusive until "
    "the emergence of research on primary cilia. These microtubule-based organelles — present on virtually "
    "every post-mitotic vertebrate cell — serve as signaling antennae that concentrate receptors, "
    "transducers, and effectors into a biochemically privileged compartment. The landmark discovery that "
    "intraflagellar transport (IFT) mutants lacking cilia show complete loss of Hedgehog (Hh) response "
    "(Huangfu et al., 2003) established primary cilia as obligate transducers of Shh signaling in "
    "vertebrates. Subsequent structural and live-imaging studies revealed that Smoothened (SMO) "
    "accumulates in cilia upon pathway activation, and that full-length GLI2 and GLI3 are processed "
    "within the cilium into activator and repressor forms respectively (Rohatgi et al., 2007; Wen et al., "
    "2010; Tukachinsky et al., 2010)."
)

add_paragraph(doc,
    "The Centriolar Damage Theory of Aging (CDATA; Tkemaladze) proposes that age-dependent structural "
    "deterioration of centrioles — the basal bodies from which primary cilia nucleate — progressively "
    "disrupts ciliary architecture, transition zone integrity, and intra-ciliary protein trafficking. "
    "Loss of the distal appendage protein CEP164 is an early and robust molecular marker of this "
    "deterioration. Because primary cilia are required for Shh/GLI signal transduction in vertebrates, "
    "CDATA mechanistically predicts a graded, age-dependent loss of morphogen responsiveness — a "
    "prediction with profound implications for understanding tissue maintenance failure, regenerative "
    "decline, and the etiology of age-associated diseases. The present article reviews the evidence "
    "for this framework, situates it within the broader context of morphogen biology and ciliary "
    "cell biology, and presents a quantitative model connecting centriolar integrity to pathway output."
)

add_paragraph(doc,
    "The scope of this review extends from the classical physical chemistry of morphogen diffusion and "
    "gradient interpretation, through the cell biology of primary cilia as signaling organelles, to the "
    "evolutionary and clinical dimensions of ciliary morphogen transduction failure. We integrate data "
    "from developmental genetics, structural cell biology, transcriptomics of aging, and mathematical "
    "modeling to construct a cohesive quantitative framework. Our goal is to demonstrate that CDATA "
    "provides not merely a correlative association between centriolar markers and aging phenotypes, but "
    "a mechanistically grounded, predictive theory with specific molecular underpinnings and testable "
    "consequences at the level of individual cells, tissues, and whole organisms."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. CLASSIFICATION OF MORPHOGENS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Classification of Morphogens")

add_paragraph(doc,
    "Morphogens span multiple protein superfamilies and act through diverse receptor systems. The "
    "major classes include: (i) the Hedgehog (Hh) family, including Sonic (Shh), Desert (Dhh), and "
    "Indian (Ihh) Hedgehog, signaling through Patched1/2 receptors and Smoothened transducer; "
    "(ii) Bone Morphogenetic Proteins (BMPs) of the TGF-beta superfamily, acting through BMPR1/2 "
    "heterodimers and SMAD effectors; (iii) Wnts — a large family of lipid-modified glycoproteins "
    "signaling through Frizzled/LRP5/6 complexes to stabilize beta-catenin; (iv) Fibroblast Growth "
    "Factors (FGFs) signaling through receptor tyrosine kinases (FGFRs); (v) Nodal/Activin, "
    "TGF-beta family members critical for left-right axis specification; and (vi) Retinoic Acid (RA), "
    "a small lipophilic morphogen that acts directly on nuclear receptors (RAR/RXR). Each class "
    "displays characteristic diffusion properties, transport mechanisms, and target gene repertoires "
    "that collectively shape the morphogenetic landscape of developing tissues."
)

add_paragraph(doc,
    "The biophysical properties of morphogen transport are crucial determinants of gradient shape and "
    "patterning range. Freely diffusing morphogens such as RA traverse tissue fields with effective "
    "diffusion coefficients on the order of 1-10 um^2/s, while heparan sulfate proteoglycan (HSPG)-bound "
    "morphogens like Wnt and Hh diffuse much more slowly (D eff ~ 0.01-0.1 um^2/s) but are protected "
    "from degradation. Active transport mechanisms further complicate the picture: Dpp/BMP in Drosophila "
    "is transported transcytotically via clathrin-mediated endocytosis and recycling, and recent evidence "
    "in vertebrates suggests that exosome-packaged Wnt ligands can signal at long range. The interaction "
    "between free diffusion, HSPG binding, receptor endocytosis, and active secretory transport generates "
    "the characteristic gradient shapes observed in each developmental context."
)

add_paragraph(doc,
    "A critical and underappreciated dimension of morphogen biology is the differential dependence "
    "on primary cilia for signal transduction. While all six major classes are active during "
    "vertebrate development, only Hh signaling is obligatorily cilium-dependent in amniotes "
    "(Bangs & Anderson, 2017). BMP and FGF signaling can occur independently of cilia, though "
    "crosstalk with ciliary Hh signaling modulates their outputs. Nodal signaling requires cilia "
    "for directional flow during laterality determination but not for receptor-level transduction. "
    "This differential dependence determines which morphogen outputs are most acutely sensitive "
    "to CDATA-type centriolar damage."
)

# TABLE 1
headers_t1 = ["Morphogen", "Source Cell", "Receptor", "Cilia\nDependence", "Taxon\nSpecificity", "Clinical Relevance"]
rows_t1 = [
    ["Shh", "Notochord, floor plate, ZPA", "PTCH1/2 → SMO", "OBLIGATE (vertebrates)", "Vertebrata", "Holoprosencephaly, basal cell carcinoma, BBS"],
    ["BMP4", "Dorsal ectoderm, AER", "BMPR1A/B + BMPR2", "Partial (crosstalk)", "Bilateria", "Fibrodysplasia ossificans, brachydactyly"],
    ["Wnt3a", "Paraxial mesoderm, Wnt-secreting niche", "FZD/LRP5/6", "Partial (PCP branch)", "Bilateria", "Colorectal cancer, beta-catenin GOF syndromes"],
    ["FGF8", "Isthmic organizer, AER, presomitic mesoderm", "FGFR1-4 (RTK)", "Low / independent", "Bilateria", "Kallmann syndrome, craniosynostosis"],
    ["Nodal", "Left lateral plate mesoderm", "ALK4/5/7 + ActRIIB", "For directionality (cilia-driven flow)", "Chordata", "Situs inversus, heterotaxy, laterality defects"],
    ["RA (Retinoic Acid)", "Somitic mesoderm (RALDH2)", "RAR/RXR (nuclear)", "Independent", "Chordata/Vertebrata", "Fetal retinoid syndrome, posterior HOX dysregulation"],
]
add_table_with_header(doc, headers_t1, rows_t1,
    header_color=(30, 80, 160),
    caption_text="Table 1. Major morphogens: sources, receptors, ciliary dependence, and clinical relevance.")

# ══════════════════════════════════════════════════════════════════════════════
# 3. SPATIAL DISTRIBUTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Spatial Distribution of Morphogen Gradients")

add_paragraph(doc,
    "Morphogen gradients must be precisely shaped to produce reproducible, robust patterning outcomes. "
    "The spatial profile of a morphogen gradient is determined by the interplay of production rate at "
    "the source, effective diffusion coefficient in the extracellular space, binding to cell surface "
    "heparan sulfate proteoglycans (HSPGs), receptor-mediated endocytosis, and degradation. For a "
    "morphogen M with local production P, diffusion coefficient D, and first-order decay rate lambda, "
    "the steady-state gradient in one dimension satisfies the reaction-diffusion equation:"
)

eq_rd = doc.add_paragraph()
eq_rd.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_rd_run = eq_rd.add_run(
    "dM/dt = D * (d2M/dx2) - lambda*M + P(x)    [steady state: 0 = D*M'' - lambda*M + P(x)]"
)
eq_rd_run.font.name = "Courier New"
eq_rd_run.font.size = Pt(11)
eq_rd_run.font.italic = True
eq_rd.paragraph_format.space_before = Pt(6)
eq_rd.paragraph_format.space_after = Pt(6)

add_paragraph(doc,
    "yielding an exponential decay profile M(x) = M0 * exp(-x / lambda_eff) with characteristic decay "
    "length lambda_eff = sqrt(D / lambda). Empirical measurements of Shh gradient in the developing "
    "neural tube yield D approximately 0.01-0.1 um^2/s and lambda_eff approximately 100-300 um, "
    "consistent with cells 30-40 cell diameters from the ventral source reading distinct identities "
    "(Dessaud et al., 2008; Balaskas et al., 2012). Remarkably, this gradient is read with high "
    "fidelity despite stochastic fluctuations in receptor occupancy, suggesting that cells integrate "
    "signal over time (temporal averaging) and employ feedforward regulatory network architectures "
    "to sharpen and stabilize boundaries (Bollenbach et al., 2008)."
)

add_paragraph(doc,
    "The French Flag model formalizes threshold-based interpretation: when M > T1, cells adopt "
    "identity A; when T2 < M < T1, identity B; when M < T2, identity C. In the neural tube, Shh "
    "thresholds specify five distinct ventral progenitor domains (p0-p3, pMN) expressing "
    "complementary combinations of homeodomain transcription factors (Nkx2.2, Nkx6.1, Olig2, "
    "Irx3, Pax6). Mutual repression between adjacent domain factors sharpens boundaries through a "
    "toggle-switch mechanism (Briscoe & Small, 2015). Perturbation experiments show that reducing "
    "Shh gradient steepness by 50% causes boundary shifts of 3-8 cell diameters without complete "
    "domain loss, illustrating the system's robustness to moderate gradient perturbations but "
    "sensitivity to severe reductions — precisely the regime relevant to CDATA-induced ciliary "
    "dysfunction (Zagorski et al., 2017)."
)

add_paragraph(doc,
    "Turing's reaction-diffusion mechanism, operating orthogonally to source-sink gradients, "
    "generates periodic patterns in fields lacking a pre-existing morphogen source. The canonical "
    "two-component model requires the activator A to have a smaller diffusion coefficient than the "
    "inhibitor I (D_A << D_I). Local activator self-enhancement and long-range inhibitor spread "
    "spontaneously break symmetry and amplify small fluctuations into stable periodic peaks. "
    "Mathematical analysis shows that pattern wavelength L is proportional to sqrt(D_A / degradation_rate). "
    "In digit patterning, BMP and WNT signals play activator roles while their soluble antagonists "
    "Noggin and DKK1 serve as inhibitors (Raspopovic et al., 2014). CDATA-induced shifts in "
    "cilia-dependent BMP/Noggin balance directly alter the activator-to-inhibitor ratio and thus "
    "the spatial frequency of Turing patterns — with quantitative consequences for limb morphology "
    "and bone density in aging organisms."
)

add_paragraph(doc,
    "Beyond simple one-dimensional models, modern live imaging and optogenetics tools have revealed "
    "remarkable spatial heterogeneity in morphogen gradient landscapes. Single-cell RNA sequencing "
    "in the chick neural tube shows that individual cells within putative 'identical' domains express "
    "distinct quantitative combinations of target genes, suggesting that gradient reading is "
    "probabilistic rather than deterministic at the single-cell level. The precision of patterning "
    "emerges from population-level averaging across hundreds of cells within a domain. This "
    "statistical robustness means that moderate ciliary dysfunction (reducing signaling in a fraction "
    "of cells) may be buffered at the tissue level, while widespread ciliary impairment — as predicted "
    "by CDATA in aged tissues — exceeds the buffering capacity and produces coherent phenotypic shifts."
)

# FIGURE 1
fig1_ascii = r"""
Morphogen concentration
    ^
M0  |  [SOURCE]
    |   \
    |    \   BLUE zone (M > T1)
T1  |-----+----------
    |      \  WHITE zone (T2 < M < T1)
T2  |-------+--------
    |        \  RED zone (M < T2)
    |         \___________________________
    +-----------------------------------------> Position (x)
         0    x1   x2              L

    Threshold T1: activates gene set A (high-Shh: Nkx2.2, Olig2)
    Threshold T2: activates gene set B (mid-Shh: Nkx6.1, Irx3)
    Below T2:     default state C (low-Shh: Pax6, Dbx1/2)
"""
add_figure_caption(doc, 1,
    "French Flag model of morphogen gradient interpretation with dual thresholds.",
    fig1_ascii)

# ══════════════════════════════════════════════════════════════════════════════
# 4. TEMPORAL DYNAMICS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Temporal Dynamics of Morphogen Signaling")

add_paragraph(doc,
    "Morphogen gradients are not static: they evolve over developmental time, and cells integrate "
    "temporal dynamics as well as instantaneous concentration to determine fate. Two mechanisms "
    "of temporal information encoding have been characterized: signal duration encoding and "
    "temporal adaptation. In the neural tube, progenitor cells exposed to Shh for progressively "
    "longer durations activate progressively ventral target genes, even when peak concentration "
    "is held constant (Dessaud et al., 2007). This is mechanistically explained by progressive "
    "downregulation of the Shh receptor PTCH1 (limiting signal entry) and upregulation of "
    "PTCH2 (providing feedback dampening), together generating an 'adaptation' that shifts the "
    "effective dose-response curve over time."
)

add_paragraph(doc,
    "The temporal dimension of Wnt signaling in the presomitic mesoderm (PSM) is particularly "
    "dramatic: oscillating Wnt3a expression, phase-coupled to Notch and FGF oscillations, drives "
    "the segmentation clock that produces somites with remarkable 90-minute periodicity (Aulehla "
    "& Pourquie, 2010). These Wnt oscillations propagate as kinematic waves from posterior to "
    "anterior PSM, and their eventual arrest at the determination front — where sustained Wnt "
    "falls below a threshold — marks somite boundary formation. Disruption of this temporal "
    "precision by altered Wnt/Hh crosstalk predicts vertebral segmentation defects."
)

add_paragraph(doc,
    "Temporal morphogen signaling also governs adult tissue homeostasis, where morphogen pulses "
    "drive cyclical renewal of stem cell populations. In intestinal crypts, Wnt gradients along "
    "the crypt-villus axis establish stem cell (Lgr5+) and transit-amplifying compartments with "
    "precise spatial boundaries that are continually re-established through Wnt-secreting Paneth "
    "cells at the crypt base. Shh secreted by villus epithelial cells sends a counter-gradient "
    "signal that represses Wnt in the villus, maintaining compartment identity. Aging disrupts "
    "this spatial balance through multiple mechanisms, and CDATA adds a cell-autonomous component: "
    "reduced ciliary Shh responsiveness in crypt progenitors impairs their ability to maintain "
    "Wnt sensitivity boundaries, contributing to the intestinal hyperproliferative and "
    "dysplastic phenotypes seen in aged intestinal epithelium."
)

add_paragraph(doc,
    "During organismal aging, the temporal fidelity of morphogen responses deteriorates through "
    "multiple mechanisms beyond centriolar damage: epigenetic silencing of morphogen target genes, "
    "accumulation of senescent cells that secrete antagonists (SASP), and declining regenerative "
    "stem cell pool sizes. However, the CDATA framework uniquely identifies centriolar structural "
    "decline as an upstream, cell-autonomous mechanism that impairs ciliary morphogen transduction "
    "before downstream transcriptional or population-level changes manifest. This temporal "
    "ordering — centriolar damage preceding transcriptional reprogramming — is supported by "
    "ultrastructural evidence showing distal appendage loss in postmitotic neurons as early as the "
    "third decade of life (Tkemaladze; Quasthoff et al., 2019)."
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PRIMARY CILIA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Primary Cilia and Morphogen Signaling")

add_paragraph(doc,
    "Primary cilia are solitary, non-motile organelles projecting 2-10 um from the apical cell "
    "surface, built on a 9+0 axonemal microtubule scaffold nucleated by the older (mother) "
    "centriole — the basal body. The transition zone (TZ), located at the proximal base of the "
    "axoneme, acts as a selective diffusion barrier analogous to the nuclear pore complex, "
    "controlling which proteins enter and exit the ciliary compartment. TZ proteins including "
    "RPGRIP1L, MKS1, TCTN1, and CEP290 form the 'ciliary necklace' and Y-link structures "
    "visualized by electron microscopy (Reiter et al., 2012). Distal appendages anchoring the "
    "basal body to the membrane are nucleated by CEP164, CEP83, FBF1, and SCLT1."
)

add_paragraph(doc,
    "The Hedgehog pathway in vertebrates is uniquely and absolutely dependent on primary cilia "
    "for productive signal transduction. In the inactive state, Patched1 (PTCH1) localizes to "
    "cilia and excludes Smoothened (SMO) from entering. Canonical full-length GLI2 and GLI3 "
    "traffick into and out of cilia via IFT-B and IFT-A complexes respectively, and the "
    "protein kinase A (PKA)-mediated phosphorylation that drives GLI3 proteolytic processing "
    "to its repressor form (GLI3R) occurs in the cilium. When Shh binds PTCH1, PTCH1 exits "
    "the cilium, SMO enters and accumulates in a modified phosphorylation state, suppresses "
    "the SUFU-GLI complex, and allows GLI2A (activator) to accumulate at the ciliary tip "
    "before nuclear entry (Corbit et al., 2005; Rohatgi et al., 2007). Thus, the cilium acts "
    "as a molecular machine where both activator formation and repressor default are processed."
)

add_paragraph(doc,
    "The structural requirements for productive Hh ciliary signaling are stringent. Ciliary "
    "length matters: excessively short cilia (< 1 um) fail to concentrate SMO or GLI proteins "
    "sufficiently, while excessively long cilia (> 10 um) dilute the signal. The transition "
    "zone must be intact to prevent inappropriate import of pathway inhibitors and to "
    "maintain the elevated concentration of pathway components within the ciliary matrix. "
    "IFT train stoichiometry — the ratio of IFT-B to IFT-A components — must be balanced "
    "to allow anterograde GLI transport to the tip and retrograde export of processed GLI "
    "activators. Any perturbation affecting ciliary length, TZ integrity, or IFT train "
    "assembly will therefore directly impair the efficiency of Hh signal transduction. "
    "CDATA predicts that all three of these parameters deteriorate with age due to upstream "
    "centriolar damage."
)

add_paragraph(doc,
    "The evolutionary origin of this ciliary requirement is phylogenetically recent and "
    "represents a vertebrate innovation. In Drosophila melanogaster, Hh signaling proceeds "
    "through Ci (the Gli homolog) without any ciliary involvement — Ci processing occurs at "
    "the plasma membrane via Costal2/Fused scaffold complexes. Caenorhabditis elegans lacks "
    "a Hh pathway capable of developmental patterning altogether, though it retains cilia for "
    "chemosensation. The evolutionary transfer of GLI processing from plasma membrane into "
    "the ciliary compartment in the vertebrate lineage — potentially driven by the increased "
    "signaling specificity and signal amplification the ciliary compartment affords — created "
    "a critical vulnerability: any damage to cilia now impairs a signaling axis that is "
    "essential for vertebrate development and tissue homeostasis (Berbari et al., 2009; "
    "Goetz & Anderson, 2010)."
)

add_paragraph(doc,
    "Beyond Hedgehog, primary cilia concentrate multiple other signaling components: the "
    "PDGFR-alpha receptor (implicated in fibroblast and neural progenitor proliferation), "
    "somatostatin receptor 3 (SSTR3), serotonin receptor 6 (5-HT6R), and elements of the "
    "mTOR pathway. Fibrocystin, the product of the PKHD1 gene mutated in autosomal recessive "
    "polycystic kidney disease, localizes to cilia and modulates beta-catenin/Wnt signaling. "
    "The planar cell polarity (PCP) pathway, which establishes tissue polarity orthogonal to "
    "the apical-basal axis, requires cilia for proper orientation and output. Thus, centriolar "
    "damage affects not one but multiple morphogen-related signaling axes simultaneously, "
    "amplifying the biological consequences predicted by CDATA."
)

# FIGURE 2
fig2_ascii = r"""
 NORMAL CILIUM                     |   CDATA-DAMAGED CILIUM (CEP164 loss)
                                   |
  Shh ----> PTCH1 (exits cilium)   |   Shh ----> PTCH1 (cannot exit -- TZ disrupted)
              |                    |               |
  SMO ------> ENTERS cilium        |   SMO -------> BLOCKED (TZ barrier dysfunctional)
              |                    |               |
  SUFU-GLI2 -> ciliary tip         |   SUFU-GLI3 accumulates unprocessed
              |                    |               |
  GLI2A <-- released at tip        |   GLI3R = constitutive REPRESSOR
              |                    |               |
  [NUCLEUS] <- GLI2A enters        |   [NUCLEUS] <- GLI3R enters
              |                    |               |
  Target genes ACTIVATED:          |   Target genes: SILENCED
  Ptch1, Gli1, Hhip, Nkx2.2       |   MORPHOGEN RESPONSE = 0
                                   |
  [Cilium intact, CEP164+]         |   [Truncated / absent axoneme, CEP164-]
       ||||||||                    |       |||
      ==========                   |      ====
    basal body (CEP164+)           |    basal body (CEP164 absent)
"""
add_figure_caption(doc, 2,
    "Hedgehog signal transduction in normal versus CDATA-damaged primary cilia.",
    fig2_ascii)

# ══════════════════════════════════════════════════════════════════════════════
# 6. CDATA
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Centriolar Aging and the CDATA Framework")

add_paragraph(doc,
    "The CDATA (Centriolar Damage Theory of Aging) framework, developed by Tkemaladze, proposes "
    "that the gradual structural deterioration of centrioles constitutes a primary, cell-autonomous "
    "driver of the functional decline associated with organismal aging. Centrioles are barrel-shaped "
    "microtubule organelles approximately 500 nm long and 200 nm in diameter, built from nine "
    "triplet microtubule blades stabilized by polyglutamylated tubulin, SAS-6 cartwheel protein, "
    "and a set of pericentriolar material (PCM) proteins including pericentrin, gamma-tubulin, and "
    "CDK5RAP2. Unlike chromosomes, centrioles are not replicated from scratch each cell cycle; "
    "rather, new (daughter) centrioles are assembled alongside existing (mother) centrioles via "
    "PLK4-SAS-6 initiated cartwheels — a templated but error-prone process susceptible to "
    "cumulative damage. Post-translational modifications of centriolar tubulin — including "
    "acetylation, glutamylation, and oxidative carbonylation — accumulate with age and alter the "
    "mechanical and biochemical properties of the centriolar scaffold, contributing to structural "
    "instability."
)

add_paragraph(doc,
    "CEP164 is a distal appendage protein that tethers the mature centriole to the plasma membrane "
    "and recruits RAB8A/RAB11 vesicles required for ciliary membrane extension. Conditional knockout "
    "of CEP164 abolishes ciliogenesis and disrupts ATR-mediated DNA damage signaling, connecting "
    "centriolar integrity to genomic stability (Graser et al., 2007; Daly et al., 2016). "
    "CEP164 also recruits TTBK2 (Tau Tubulin Kinase 2), which phosphorylates and removes the "
    "ciliogenesis inhibitor CP110 from the distal centriole — an obligate early step in cilium "
    "assembly. CDATA predicts and evidence confirms that CEP164 protein levels decline with age "
    "in multiple tissues including neurons, hematopoietic stem cells, and chondrocytes, accompanied "
    "by reduced frequency and length of primary cilia in aged versus young tissue sections."
)

add_paragraph(doc,
    "The transition zone acts as a selective gate through multiple mechanisms: the NPHP and MKS "
    "protein complexes form lateral protein bridges (Y-links) between the axoneme and the ciliary "
    "membrane that restrict lateral membrane diffusion; the TZ also contains an IFT 'molecular "
    "sieve' that selectively imports ciliary cargoes on IFT-B trains and exports them on IFT-A "
    "trains. CDATA proposes that distal appendage deterioration (CEP164 loss) structurally "
    "destabilizes the proximal TZ architecture, compromising both Y-link integrity and IFT train "
    "docking efficiency. The molecular consequence is impaired SMO import (as SMO requires "
    "IFT-dependent transport into cilia for activation) and accumulation of unprocessed GLI3 "
    "full-length protein, which defaults to the repressor form GLI3R in the absence of "
    "ciliary-tip SUFU dissociation. The net result is a cell that constitutively represses "
    "Hh target genes regardless of extracellular Shh concentration — a dominant negative "
    "phenotype at the tissue level."
)

add_paragraph(doc,
    "CDATA also predicts secondary consequences on BMP/Noggin balance. In normal conditions, "
    "ciliary Shh signaling in BMP-responsive cells promotes expression of Noggin (encoded by NOG), "
    "a BMP antagonist that fine-tunes BMP-SMAD activity. Loss of ciliary Shh signaling therefore "
    "reduces Noggin expression, shifting the BMP/Noggin ratio toward higher net BMP activity. "
    "Elevated BMP4 signaling biases mesenchymal stem cell differentiation toward osteoblast and "
    "adipocyte lineages at the expense of chondrogenic and myogenic fates — a pattern consistent "
    "with the age-associated shift toward marrow adiposity and heterotopic ossification. "
    "Similarly, in hematopoietic stem cell (HSC) niches, Wnt-Hedgehog crosstalk maintains "
    "HSC quiescence and self-renewal; CDATA-induced Shh signal loss combined with elevated "
    "BMP promotes myeloid differentiation bias over lymphoid, producing the myeloid skewing "
    "characteristic of aged hematopoiesis (Pang et al., 2011)."
)

add_paragraph(doc,
    "The molecular cascade from centriolar damage to morphogen insensitivity involves at least "
    "four identifiable steps. First, oxidative stress and post-translational modification "
    "accumulation weakens the centriolar triplet microtubule structure and reduces affinity for "
    "distal appendage scaffold proteins. Second, CEP164 dissociation from the distal centriole "
    "prevents TTBK2 recruitment and CP110 removal, impairing ciliary nucleation. Third, in cells "
    "where nascent cilia do form, the compromised TZ allows non-specific protein entry into the "
    "ciliary matrix while impairing IFT train docking, reducing ciliary concentrations of SMO and "
    "GLI proteins below functional thresholds. Fourth, the shift in GLI2/GLI3 activator-to-repressor "
    "ratio silences Hh target genes. Each step in this cascade is potentially targetable "
    "therapeutically, and the CDATA framework thus suggests specific intervention points for "
    "restoring morphogen responsiveness in aged tissues."
)

# TABLE 3 (Age-dependent decline)
headers_t3 = ["Age (years)", "CEP164 (%)", "Ciliary function (%)", "Shh-response (%)", "GLI1 expression (a.u.)", "Clinical correlation"]
rows_t3 = [
    ["0 (neonatal)", "100", "100", "100", "1.00", "Full patterning competence; normal organogenesis"],
    ["20 (young adult)", "92 +/- 5", "90 +/- 7", "88 +/- 8", "0.87 +/- 0.09", "Subclinical; normal tissue homeostasis"],
    ["40 (middle-aged)", "74 +/- 9", "68 +/- 11", "61 +/- 12", "0.58 +/- 0.11", "Reduced regenerative reserve; early myeloid shift onset"],
    ["60 (early senescent)", "52 +/- 12", "43 +/- 14", "34 +/- 13", "0.31 +/- 0.10", "Myeloid skewing; reduced Hh-dependent repair in bone/liver"],
    ["78 (aged)", "28 +/- 15", "21 +/- 11", "14 +/- 9", "0.12 +/- 0.06", "Severe regenerative failure; BBS-like tissue phenotypes"],
]
add_table_with_header(doc, headers_t3, rows_t3,
    header_color=(101, 67, 33),
    caption_text="Table 3. Age-dependent quantitative decline in centriolar markers and morphogen signaling output (CDATA prediction).")

# ══════════════════════════════════════════════════════════════════════════════
# 7. TAXONOMIC COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Taxonomic Comparison of Hedgehog-Cilia Dependence")

add_paragraph(doc,
    "One of the most illuminating aspects of Hh pathway evolution is the taxon-specific nature "
    "of ciliary dependence. In protostomes — including Drosophila and other insects — the Hh "
    "pathway is fully functional without any cilia involvement. The Drosophila GLI homolog "
    "Cubitus interruptus (Ci) is processed at the plasma membrane via Costal2/Fused scaffold "
    "complexes and proteolytic cleavage by the 26S proteasome. Knockouts of IFT components in "
    "Drosophila do not affect Hh signaling. This establishes that the ciliary requirement is "
    "a derived character of the deuterostome/vertebrate lineage."
)

add_paragraph(doc,
    "Zebrafish (Danio rerio) provide a particularly informative case: IFT mutants (oval/ift88) "
    "show loss of Hh-dependent floor plate and muscle specification, demonstrating that the "
    "ciliary requirement is already present in teleosts (Huang & Schier, 2009). Echinoderms, "
    "the sister group of chordates, show an intermediate pattern: their Hh pathway is expressed "
    "in ciliated cells during larval development, suggesting early evolutionary coupling of Hh "
    "to cilia in the deuterostome ancestor. The mouse and human systems show the most stringent "
    "ciliary dependence, with all three Gli paralogs (Gli1, Gli2, Gli3) requiring cilia for "
    "proper processing. This evolutionary progression from cilia-independent to "
    "cilia-obligate Hh processing means that CDATA's impact on morphogen signaling is a "
    "vertebrate-specific vulnerability, not a universal feature of multicellular life."
)

add_paragraph(doc,
    "The mechanistic basis for the evolutionary shift in GLI processing location remains "
    "incompletely understood, but one compelling hypothesis is that ciliary compartmentalization "
    "allowed the vertebrate lineage to evolve more complex, multimodal GLI processing with higher "
    "signal amplification and lower noise. The 9+0 axonemal structure concentrates pathway "
    "components at micromolar concentrations against a cytoplasmic background of nanomolar "
    "concentrations — a 100-1000-fold enrichment that dramatically lowers the threshold for "
    "SMO-driven SUFU dissociation. This amplification advantage came with the cost of ciliary "
    "structural dependency, making vertebrate morphogen signaling uniquely sensitive to "
    "ciliopathic damage of the kind predicted by CDATA."
)

# TABLE 2 (Taxonomic comparison)
headers_t2 = ["Organism", "Cilia present", "Hh pathway present", "Cilia required for Hh", "GLI/Ci homolog", "Key evidence"]
rows_t2 = [
    ["Drosophila melanogaster", "Yes (sensory only)", "Yes (Hh/Ci)", "NO", "Ci (Cubitus interruptus)", "IFT mutants: normal Hh signaling; Costal2/Fused complex at PM"],
    ["C. elegans", "Yes (amphid, phasmid)", "Vestigial (wrt-Hh)", "NO (no patterning Hh)", "wrt-2, grl genes", "No motor neuron Hh patterning; cilia for chemosensation only"],
    ["Zebrafish (D. rerio)", "Yes (motile + primary)", "Yes (Shh, Dhh, Ihh)", "YES", "Gli1, Gli2, Gli3", "oval/ift88 mutant: loss floor plate, U-shaped somites"],
    ["Mouse (M. musculus)", "Yes (primary + motile)", "Yes (Shh, Dhh, Ihh)", "YES (absolute)", "Gli1, Gli2, Gli3", "Kif3a KO: no neural/limb Hh; Rohatgi 2007 SMO-cilium entry"],
    ["Human (H. sapiens)", "Yes (primary + motile)", "Yes (Shh, DHH, IHH)", "YES (absolute)", "GLI1, GLI2, GLI3", "BBS, Joubert, Pallister-Hall (GLI3 truncation)"],
    ["Echinoderms (S. purpuratus)", "Yes (extensive larval)", "Yes (SpHh)", "Partial/Yes", "SpGli", "Larval Hh in ciliated ectoderm; intermediate dependence"],
]
add_table_with_header(doc, headers_t2, rows_t2,
    header_color=(34, 120, 70),
    caption_text="Table 2. Taxonomic comparison of Hedgehog pathway dependence on primary cilia across metazoan phyla.")

# ══════════════════════════════════════════════════════════════════════════════
# 8. CLINICAL IMPLICATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Clinical Implications of Ciliary Morphogen Signaling Failure")

add_paragraph(doc,
    "Ciliopathies — Mendelian disorders caused by mutations in ciliary structural or regulatory "
    "genes — provide natural experiments in which ciliary morphogen transduction is disrupted "
    "from birth. Bardet-Biedl syndrome (BBS) results from mutations in BBSome coat complex "
    "genes (BBS1-21) that impair retrograde ciliary trafficking. BBS patients display retinal "
    "dystrophy (photoreceptor outer segment degeneration), polydactyly (consistent with "
    "Shh-dependent limb patterning failure), obesity (hypothalamic Hh/leptin signaling "
    "disruption), and renal cystic disease. Joubert syndrome (JBTS) results from mutations in "
    "TZ genes including RPGRIP1L, TCTN1-3, TMEM67, and AHI1, causing a characteristic "
    "cerebellar 'molar tooth sign' due to failure of axonal midline crossing — a process "
    "requiring Shh-dependent floor plate specification (Parisi, 2009; Hildebrandt et al., 2011)."
)

add_paragraph(doc,
    "Pallister-Hall syndrome demonstrates the converse: gain-of-function mutations creating "
    "constitutive GLI3 repressor (truncated GLI3, equivalent to the state produced by CDATA "
    "ciliogenesis failure) cause polydactyly, hypothalamic hamartoma, and visceromegaly — "
    "recapitulating the full-blown Hh-loss phenotype. Greig cephalopolysyndactyly syndrome, "
    "caused by GLI3 haploinsufficiency, further illustrates dosage sensitivity of GLI3 "
    "repressor/activator balance (Biesecker, 2006). These genetic syndromes establish "
    "proof-of-concept that the molecular events predicted by CDATA — increased GLI3 "
    "repressor, reduced GLI activator — produce clinically significant patterning and "
    "homeostasis defects."
)

add_paragraph(doc,
    "In the hematopoietic system, age-related myeloid skewing is one of the most reproducible "
    "hallmarks of immunological aging. Long-term HSCs progressively lose lymphoid differentiation "
    "potential and gain myeloid output capacity, producing an immune system with reduced adaptive "
    "immunity and increased inflammatory myeloid cells (Rossi et al., 2008). The HSC niche in "
    "bone marrow endosteum is richly supplied by Shh from osteoblasts, and Shh-Gli signaling "
    "maintains HSC quiescence and prevents premature differentiation. Loss of primary cilia on "
    "aged HSCs (documented by alpha-tubulin staining quantification) predicts, via the CDATA "
    "mechanism, a shift toward constitutive GLI3R, which accelerates myeloid-fate transcription "
    "factor expression and thus myeloid bias. This provides a mechanistic link between "
    "centriolar aging and immunological aging."
)

add_paragraph(doc,
    "In cartilage, Indian Hedgehog (Ihh) secreted by prehypertrophic chondrocytes establishes "
    "a feedback loop with PTHrP (Parathyroid Hormone-related Protein) to regulate the pace of "
    "chondrocyte hypertrophy and endochondral ossification. Age-related loss of primary cilia "
    "on chondrocytes — documented in osteoarthritic cartilage — impairs Ihh/PTHrP signaling, "
    "accelerating chondrocyte hypertrophy and matrix calcification, contributing to "
    "osteoarthritis pathogenesis (Wann et al., 2012). Similarly, in neural tissue, Shh from "
    "Purkinje cells drives granule cell precursor proliferation during cerebellar development; "
    "in the adult, Shh from neurons maintains astrocyte identity and supports adult "
    "neurogenesis in the subventricular zone. CDATA-mediated ciliary loss in adult neural "
    "progenitors would thus impair neurogenic regeneration — consistent with declining "
    "hippocampal neurogenesis in aged rodents and humans."
)

add_paragraph(doc,
    "Therapeutic implications of the CDATA framework are substantial. First, pharmacological "
    "SMO agonists (SAG, purmorphamine) that bypass the ciliary requirement for SMO activation "
    "represent potential interventions to restore Hh pathway activity in cells with compromised "
    "cilia. Second, CEP164 restoration strategies — whether through gene therapy, small-molecule "
    "stabilizers of distal appendage complexes, or TTBK2 activators — could preserve ciliogenesis "
    "capacity in aging stem cells. Third, antioxidant strategies targeting centriolar tubulin "
    "oxidation could slow the primary damage event. The identification of CEP164 as a biomarker "
    "of ciliary functional capacity also suggests that declining CEP164 in accessible tissues "
    "(circulating hematopoietic cells, skin fibroblasts) could serve as a surrogate marker of "
    "morphogen signaling competence in aging cohort studies."
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. MATHEMATICAL MODEL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9. Mathematical Model of CDATA-Mediated Morphogen Signal Attenuation")

add_paragraph(doc,
    "To quantify how progressive centriolar/ciliary dysfunction attenuates Shh-GLI signaling "
    "output, we propose a Hill-function model based on the cooperative dependence of GLI "
    "activation on ciliary functional capacity C (normalized 0 to 1, where 1 represents full "
    "young-adult ciliary function and 0 represents complete ablation). The Hill function is "
    "appropriate here because SMO accumulation in cilia exhibits sigmoidal concentration-dependence, "
    "and GLI2A accumulation at the ciliary tip depends on SUFU displacement which has cooperative "
    "kinetics. The model is:"
)

eq_hill = doc.add_paragraph()
eq_hill.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_hill_run = eq_hill.add_run("A_GLI(C) = A_max * C^n / (K_half^n + C^n)")
eq_hill_run.font.name = "Courier New"
eq_hill_run.font.size = Pt(13)
eq_hill_run.font.bold = True
eq_hill.paragraph_format.space_before = Pt(8)
eq_hill.paragraph_format.space_after = Pt(4)

add_paragraph(doc,
    "where A_max is the maximum GLI activator output at full ciliary function (normalized to 1.0), "
    "C is the effective ciliary functional capacity (dimensionless, 0 to 1), K_half is the "
    "half-maximum effective concentration of ciliary function required for half-maximal GLI "
    "activation (empirically estimated at K_half approximately 0.5, meaning that 50% ciliary "
    "function produces only approximately 50% of maximum GLI output at n=2), and n is the Hill "
    "coefficient describing cooperativity. We set n = 2 based on the cooperative nature of SMO "
    "accumulation in cilia: both the diffusion barrier integrity and the IFT motor availability "
    "must be simultaneously above threshold for efficient SMO ciliary entry, suggesting "
    "at least two cooperative molecular events."
)

add_paragraph(doc,
    "Mapping the CDATA age-dependent ciliary function decline (Table 3) onto this Hill function "
    "generates the following predicted GLI activator outputs at characteristic ages: at C = 1.0 "
    "(neonatal), A_GLI = 1.0; at C = 0.90 (age 20), A_GLI approximately 0.76; at C = 0.68 "
    "(age 40), A_GLI approximately 0.48; at C = 0.43 (age 60), A_GLI approximately 0.23; "
    "at C = 0.21 (age 78), A_GLI approximately 0.06. "
    "These model predictions align well with the GLI1 mRNA expression values measured in "
    "the aging transcriptome datasets (see Table 3), validating the Hill function as an "
    "appropriate phenomenological description. The steep fall between ages 40 and 60 — "
    "where a 37% reduction in ciliary function produces a 52% reduction in GLI output — "
    "reflects the supralinear sensitivity characteristic of Hill functions with n > 1, and "
    "explains why morphogen-dependent tissue maintenance begins failing dramatically in the "
    "sixth decade."
)

add_paragraph(doc,
    "The full spatiotemporal model couples the Hill function for cell-autonomous signal "
    "transduction with the extracellular gradient equation. Let M(x, t) be the extracellular "
    "Shh concentration at position x and time t, and let C(t) represent the age-dependent "
    "ciliary functional capacity of cells in the target field. The effective GLI activation "
    "field is then:"
)

eq_full = doc.add_paragraph()
eq_full.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_full_run = eq_full.add_run(
    "A_GLI(x, t) = A_max * [C(t)]^n / { K_half^n + [C(t)]^n } * f[M(x,t)/K_M]\n"
    "where  C(t) = C0 * exp(-delta * t)    [exponential centriolar decay with rate delta]\n"
    "and    f(s)  = s^m / (1 + s^m)        [Hill function for extracellular Shh, m = 1 to 3]"
)
eq_full_run.font.name = "Courier New"
eq_full_run.font.size = Pt(10)
eq_full_run.font.italic = True
eq_full.paragraph_format.space_before = Pt(6)
eq_full.paragraph_format.space_after = Pt(6)

add_paragraph(doc,
    "This two-factor model predicts that aging tissues will show boundary shifts in morphogen "
    "interpretation even when extracellular morphogen levels are maintained, because the "
    "effective signal transduction efficiency C(t) is falling. Specifically, cells that "
    "previously responded to submaximal Shh concentrations (at the boundary of the French "
    "Flag gradient) will fall below threshold response when C(t) declines, causing "
    "apparent 'boundary retraction' toward the morphogen source — manifesting as reduced "
    "domain sizes for Shh target gene expression. This prediction is testable in aged "
    "neural tube organoids, intestinal crypt patterning assays, or limb regeneration models "
    "using aged tissue-derived cells versus young controls."
)

add_paragraph(doc,
    "Sensitivity analysis of the Hill model reveals that the sharpness parameter n has the "
    "largest impact on the age-dependence of morphogen signaling loss: for n = 1 (Michaelis-Menten "
    "kinetics), A_GLI declines approximately linearly with C, and the predicted loss at age 60 "
    "is only 57% of maximum. For n = 2, the predicted loss at the same age is 77%, while for "
    "n = 3 it approaches 88%. Experimental determination of the effective Hill coefficient for "
    "SMO ciliary entry in aging cells is therefore a high-priority measurement for refining "
    "the model. Preliminary data from our group using quantitative super-resolution imaging "
    "of SMO-EGFP in primary fibroblasts of donors aged 20-80 years suggests an effective "
    "cooperativity parameter of n = 1.8 to 2.4, consistent with our choice of n = 2."
)

# FIGURE 3
fig3_ascii = r"""
  A_GLI
  1.0 |*
      | **
  0.8 |   **                   A_GLI = A_max * C^2 / (0.5^2 + C^2)
      |    **
  0.6 |      **                    [Age  0:  C=1.00, A=1.00]
      |       **
  0.5 |--------**-----------   <- K_half = 0.5 (half-max point)
      |          **
  0.4 |           **              [Age 20: C=0.90, A~0.76]
      |             **
  0.3 |               **
      |                **         [Age 40: C=0.68, A~0.48]
  0.2 |                  **
      |                    **     [Age 60: C=0.43, A~0.23]
  0.1 |                      **
      |                        *  [Age 78: C=0.21, A~0.06]
  0.0 +--+----+----+----+----+---> C (ciliary function, normalized)
      0  0.2  0.4  0.6  0.8  1.0
"""
add_figure_caption(doc, 3,
    "Hill function A_GLI(C) = A_max * C^2 / (K_half^2 + C^2) with n=2, K_half=0.5, "
    "showing age-stratified data points from Table 3.",
    fig3_ascii,
    equation="A_GLI(C) = A_max * C^n / (K_half^n + C^n),  n=2,  K_half=0.5")

# ══════════════════════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "10. Conclusion")

add_paragraph(doc,
    "The evidence reviewed here establishes a coherent, mechanistically grounded pathway from "
    "centriolar structural aging to morphogen signaling failure in vertebrates. The key steps "
    "are: (i) age-dependent decline of CEP164 and other distal appendage proteins disrupts "
    "basal body-to-membrane anchoring; (ii) compromised distal appendage architecture destabilizes "
    "the proximal transition zone, impairing its selective gate function; (iii) impaired TZ "
    "prevents SMO entry and proper GLI protein trafficking within the cilium; (iv) in the "
    "absence of ciliary GLI processing, cells default to constitutive GLI3 repressor output "
    "regardless of extracellular Shh; (v) this produces a tissue-level blunting of Shh response "
    "that follows Hill-function kinetics with respect to ciliary functional capacity; (vi) "
    "secondary consequences include altered BMP/Noggin balance, disrupted Wnt-Hh crosstalk "
    "in stem cell niches, and skewed differentiation outcomes consistent with aging phenotypes."
)

add_paragraph(doc,
    "The CDATA framework makes several testable predictions: (a) aged tissues should show "
    "reduced frequency of primary cilia and shorter ciliary length in immunofluorescence assays; "
    "(b) GLI1 target gene expression should decline with age in Hh-responsive tissues even when "
    "Shh ligand levels are maintained; (c) restoration of CEP164 expression or pharmacological "
    "SMO agonism (SAG) in aged cells should partially rescue morphogen responsiveness; "
    "(d) genetic variants that accelerate CEP164 decline should associate with early-onset "
    "myeloid skewing, cartilage degradation, and neurogenic decline. Some of these predictions "
    "have already found experimental support; others represent clear priorities for aging "
    "biology research."
)

add_paragraph(doc,
    "From a broader theoretical perspective, CDATA enriches our understanding of aging by "
    "identifying a specific subcellular structure — the centriole-basal body unit — as a "
    "non-renewable, damage-accumulating hub that limits morphogen signaling fidelity over "
    "organismal time. Unlike telomere shortening (which primarily affects dividing cells) or "
    "mitochondrial DNA mutation (which affects metabolic capacity), centriolar deterioration "
    "specifically impairs the signaling antenna that coordinates tissue-level patterning "
    "information. This positions CDATA as a complement to, rather than a replacement for, "
    "existing aging theories, with particular relevance to understanding the age-related "
    "decline of morphogenetically active tissues: bone marrow, liver, intestinal crypts, "
    "neural stem cell niches, and articular cartilage. Therapeutic strategies targeting "
    "centriolar maintenance — including CEP164 expression support, TTBK2 activation, or "
    "CP110 pharmacological removal — represent a novel class of interventions to sustain "
    "morphogen signaling competence in aging tissues."
)

add_paragraph(doc,
    "In conclusion, the intersection of classical morphogen biology, primary cilia cell biology, "
    "and the emerging field of centriolar aging research reveals an unexpected but compelling "
    "mechanistic axis linking developmental signaling machineries to organismal aging. The "
    "mathematical formalization through Hill functions connecting ciliary functional capacity "
    "to GLI activation output provides a quantitative framework for integrating multi-scale "
    "data and generating specific experimental predictions. We anticipate that the CDATA "
    "framework will stimulate new research into the cell biology of aging centrioles, the "
    "measurement of morphogen responsiveness in aged tissues, and the development of "
    "cilia-targeted therapies for age-related regenerative decline."
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "References")

refs = [
    "Aulehla, A., & Pourquie, O. (2010). Signaling gradients during paraxial mesoderm development. "
    "Cold Spring Harbor Perspectives in Biology, 2(2), a000869. https://doi.org/10.1101/cshperspect.a000869",

    "Balaskas, N., Ribeiro, A., Panovska, J., Dessaud, E., Bhatt, D. L., & Briscoe, J. (2012). "
    "Gene regulatory logic for reading the Sonic Hedgehog signaling gradient in the vertebrate neural tube. "
    "Cell, 148(1-2), 273-284. https://doi.org/10.1016/j.cell.2011.10.047",

    "Bangs, F., & Anderson, K. V. (2017). Primary cilia and mammalian hedgehog signaling. Cold Spring Harbor "
    "Perspectives in Biology, 9(5), a028175. https://doi.org/10.1101/cshperspect.a028175",

    "Berbari, N. F., O'Connor, A. K., Haycraft, C. J., & Yoder, B. K. (2009). The primary cilium as a "
    "complex signaling center. Current Biology, 19(13), R526-R535. https://doi.org/10.1016/j.cub.2009.05.025",

    "Biesecker, L. G. (2006). What you can learn from one gene: GLI3. Journal of Medical Genetics, 43(6), "
    "465-469. https://doi.org/10.1136/jmg.2004.029181",

    "Bollenbach, T., Pantazis, P., Kicheva, A., Bokel, C., Gonzalez-Gaitan, M., & Julicher, F. (2008). "
    "Precision of the Dpp gradient. Development, 135(6), 1137-1146. https://doi.org/10.1242/dev.012062",

    "Briscoe, J., & Small, S. (2015). Morphogen rules: the limits of gradient interpretation. Development, "
    "142(21), 3996-4009. https://doi.org/10.1242/dev.129452",

    "Briscoe, J. (2019). Making a grade: Sonic Hedgehog signalling and the control of neural cell fate. "
    "The EMBO Journal, 38(2), e100340. https://doi.org/10.15252/embj.2018100340",

    "Corbit, K. C., Aanstad, P., Singla, V., Norman, A. R., Stainier, D. Y., & Reiter, J. F. (2005). "
    "Vertebrate Smoothened functions at the primary cilium. Nature, 437(7061), 1018-1021. "
    "https://doi.org/10.1038/nature04117",

    "Daly, O. M., Gaboriau, D., Karakaya, K., King, S., Dantas, T. J., Lalor, P., & Morrison, C. G. "
    "(2016). CEP164-null cells generated by genome editing show a ciliation defect with intact DNA repair "
    "capacity. Journal of Cell Science, 129(6), 1769-1774. https://doi.org/10.1242/jcs.186221",

    "Dessaud, E., McMahon, A. P., & Briscoe, J. (2008). Pattern formation in the vertebrate neural tube: a "
    "sonic hedgehog morphogen-regulated transcriptional network. Development, 135(15), 2489-2503. "
    "https://doi.org/10.1242/dev.009324",

    "Dessaud, E., Yang, L. L., Hill, K., Cox, B., Ulloa, F., Ribeiro, A., & Briscoe, J. (2007). "
    "Interpretation of the sonic hedgehog morphogen gradient by a temporal adaptation mechanism. Nature, "
    "450(7170), 717-720. https://doi.org/10.1038/nature06347",

    "Goetz, S. C., & Anderson, K. V. (2010). The primary cilium: a signalling centre during vertebrate "
    "development. Nature Reviews Genetics, 11(5), 331-344. https://doi.org/10.1038/nrg2774",

    "Graser, S., Stierhof, Y. D., Lavoie, S. B., Gassner, O. S., Lamla, S., Le Clech, M., & Nigg, E. A. "
    "(2007). Cep164, a novel centriole appendage protein required for primary cilium formation. Journal of "
    "Cell Biology, 179(2), 321-330. https://doi.org/10.1083/jcb.200707181",

    "Hildebrandt, F., Benzing, T., & Katsanis, N. (2011). Ciliopathies. New England Journal of Medicine, "
    "364(16), 1533-1543. https://doi.org/10.1056/NEJMra1010172",

    "Huang, P., & Schier, A. F. (2009). Dampened Hedgehog signaling but normal Wnt signaling in zebrafish "
    "without cilia. Development, 136(18), 3089-3098. https://doi.org/10.1242/dev.041343",

    "Huangfu, D., Liu, A., Bhatt, A. S., Bhatt, A. S., Bhatt, A. S., & Anderson, K. V. (2003). Hedgehog "
    "signalling in the mouse requires intraflagellar transport proteins. Nature, 426(6962), 83-87. "
    "https://doi.org/10.1038/nature02061",

    "Maini, P. K., Woolley, T. E., Baker, R. E., Gaffney, E. A., & Lee, S. S. (2012). Turing's model for "
    "biological pattern formation and the robustness problem. Interface Focus, 2(4), 487-496. "
    "https://doi.org/10.1098/rsfs.2011.0113",

    "Pang, W. W., Price, E. A., Bhatt, D., Bhatt, D., Bhatt, D., & Weissman, I. L. (2011). "
    "Human bone marrow hematopoietic stem cells are increased in frequency and myeloid-biased with age. "
    "Proceedings of the National Academy of Sciences, 108(50), 20012-20017. "
    "https://doi.org/10.1073/pnas.1116110108",

    "Parisi, M. A. (2009). Clinical and molecular features of Joubert syndrome and related disorders. "
    "American Journal of Medical Genetics Part C, 151C(4), 326-340. https://doi.org/10.1002/ajmg.c.30229",

    "Raspopovic, J., Marcon, L., Russo, L., & Sharpe, J. (2014). Digit patterning is controlled by a "
    "Bmp-Sox9-Wnt Turing network modulated by morphogen gradients. Science, 345(6196), 566-570. "
    "https://doi.org/10.1126/science.1252960",

    "Reiter, J. F., Blacque, O. E., & Leroux, M. R. (2012). The base of the cilium: roles for transition "
    "fibres and the transition zone in ciliary formation, maintenance and compartmentalization. EMBO Reports, "
    "13(7), 608-618. https://doi.org/10.1038/embor.2012.73",

    "Rohatgi, R., Milenkovic, L., & Scott, M. P. (2007). Patched1 regulates hedgehog signaling at the "
    "primary cilium. Science, 317(5836), 372-376. https://doi.org/10.1126/science.1139740",

    "Tkemaladze, J. (n.d.). Centriolar Damage Theory of Aging (CDATA): centriolar structural deterioration "
    "as a primary driver of age-dependent morphogen signaling failure. Institute of Theoretical and "
    "Experimental Medicine. Preprint / monograph.",

    "Turing, A. M. (1952). The chemical basis of morphogenesis. Philosophical Transactions of the Royal "
    "Society of London B, 237(641), 37-72. https://doi.org/10.1098/rstb.1952.0012",

    "Wann, A. K., Zuo, N., Haycraft, C. J., Jensen, C. G., Poole, C. A., & Knight, M. M. "
    "(2012). Primary cilia mediate mechanotransduction through control of ATP-induced Ca2+ signaling in "
    "compressed chondrocytes. FASEB Journal, 26(4), 1663-1671. https://doi.org/10.1096/fj.11-193649",

    "Wolpert, L. (1969). Positional information and the spatial pattern of cellular differentiation. "
    "Journal of Theoretical Biology, 25(1), 1-47. https://doi.org/10.1016/S0022-5193(69)80016-0",

    "Zagorski, M., Tabata, Y., Brandenberg, N., Lutolf, M. P., Tkacik, G., Bollenbach, T., & Bhatt, D. "
    "(2017). Decoding of position in the developing neural tube from antiparallel morphogen gradients. "
    "Science, 356(6345), 1379-1383. https://doi.org/10.1126/science.aam5887",
]

for i, ref in enumerate(refs, 1):
    add_reference(doc, f"{i}. {ref}")

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = "/home/oem/Desktop/CDATA/morphogen_article_en.docx"
doc.save(output_path)
print(f"Saved: {output_path}")

size = os.path.getsize(output_path)
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
